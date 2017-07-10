# -*- coding:utf-8 -*-
import urllib
import urllib2
import re
from docx import Document
from docx.shared import Inches

class Tool:
    replaceBR=re.compile('<br>')
    replaceST=re.compile('<strong.*?>|</strong.*?>')
    replaceImg=re.compile('<img alt="" src="')
    replaceIE=re.compile('" style=".*>')
    def replace(self,x):
        x=re.sub(self.replaceBR,"",x)
        x=re.sub(self.replaceST,"  ",x)
        x=re.sub(self.replaceImg,"",x)
        x=re.sub(self.replaceIE,"",x)
        return x.strip()

class Spider:
    def __init__(self):
        self.siteURL='http://wufazhuce.com/'
        self.tool=Tool()
        self.document = Document()
        #文章的题目
        self.title=''
#得到确定网页地址
    def getPage(self,pageIndex):
        url=self.siteURL+"article/"+str(pageIndex)
        #print url
        request=urllib2.Request(url)
        response=urllib2.urlopen(request)
        return response.read().decode('utf-8')
#读取文章标题传给全局变量self.title
    def getTitle(self,pageIndex):
        page = self.getPage(pageIndex)
        pattern = re.compile('<h2 class="articulo-titulo".*?>(.*?)</h2>',re.S)
        results=re.findall(pattern,page)
        self.title=results[0].strip()
#读取文章标题作者内容并写入TXT
    def getTxt(self, pageIndex):
        page = self.getPage(pageIndex)
        pattern = re.compile('<h2 class="articulo-titulo".*?>(.*?)</h2>.*?'
                             '<p class="articulo-autor".*?>(.*?)</p>.*?<div'
                             ' class="articulo-contenido".*?>(.*?)</div>', re.S)
        items = re.findall(pattern, page)
        for item in items:
          title=item[0]
          name=item[1]
          content=self.tool.replace(item[2])
        f = open(self.title+".txt", "w+")
        f.write(title.encode('utf-8'))
        f.write(name.encode('utf-8')+'\n')
        f.write(content.encode('utf-8'))
#读取图片URL保存在本地
    def printImg(self,imgurl):
        f = open('cc.jpg', "wb")
        u = urllib.urlopen(imgurl)
        buf = u.read()
        f.write(buf)
#读取TXT内容写入DOCX
    def getDocx(self):
        seq=re.compile('^http\w*://.+?/394',re.S)
        with open(self.title+".txt", 'r') as f:  # 打开文件
            lines = f.readlines()  # 读取所有行
            for line in lines:
                item=re.findall(seq,line)
                if item:
                    #有图片的文本
                    imgurl=item[0].strip()
                    self.printImg(imgurl)  # 将图片保存本地名字为cc.jpg
                    self.document.add_picture('cc.jpg',width=Inches(6.2))
                else:#没有图片的文本
                    line = line.replace("\n","")
                    line = line.replace("\r","")
                    line = line.decode('utf8')
                    line = line.replace(unichr(194),"")
                    line = line.replace(unichr(160),"")
                    if len(line) != 0 :
                        #print line
                        run=self.document.add_paragraph(line)
#将上述函数整理起来
    def run(self,pageIndex):
            self.getTitle(pageIndex)
            self.getTxt(pageIndex)
            self.getDocx()
            self.document.save(self.title + ".docx")

spider=Spider()
spider.run(2000)